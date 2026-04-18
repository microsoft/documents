"""Tests for section builder functions."""

from docx import Document as DocxDocument

from src.sections import (
    add_title_section,
    add_content_section,
    add_section_header_section,
    add_two_column_section,
    add_resource_box_section,
)
from src.style import Style


def test_title_section():
    doc = DocxDocument()
    style = Style()
    data = {
        "type": "title",
        "title": "My Title",
        "subtitle": "Sub line 1<br>Sub line 2",
        "notes": "Title notes",
        "image": None,
    }
    add_title_section(doc, data, style)
    # Should have paragraphs (spacing + title + subtitle + page break)
    assert len(doc.paragraphs) > 0
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "My Title" in text


def test_content_section():
    doc = DocxDocument()
    style = Style()
    data = {
        "type": "content",
        "title": "Content Heading",
        "bullets": ["Point A", "Point B", "Point C"],
        "notes": "",
        "image": None,
    }
    add_content_section(doc, data, style)
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Point A" in text
    assert "Point B" in text
    assert "Point C" in text


def test_section_header():
    doc = DocxDocument()
    style = Style()
    data = {
        "type": "section-header",
        "title": "New Section",
        "subtitle": "Breaking here",
        "notes": "",
        "image": None,
    }
    add_section_header_section(doc, data, style)
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "New Section" in text


def test_two_column():
    doc = DocxDocument()
    style = Style()
    data = {
        "type": "two-column",
        "title": "Comparison",
        "left_bullets": ["L1", "L2"],
        "right_bullets": ["R1", "R2"],
        "notes": "",
        "image": None,
    }
    add_two_column_section(doc, data, style)
    assert len(doc.tables) == 1
    table = doc.tables[0]
    assert len(table.rows) == 2
    assert "L1" in table.cell(0, 0).text
    assert "R1" in table.cell(0, 1).text


def test_resource_box():
    doc = DocxDocument()
    style = Style()
    data = {
        "type": "resource-box",
        "title": "Resources",
        "subtitle": "Links",
        "boxes": [
            {
                "label": "Docs",
                "rows": [
                    {"name": "Doc A", "url": "https://example.com/a"},
                    {"name": "Doc B", "url": "https://example.com/b"},
                ],
            }
        ],
        "notes": "",
        "image": None,
    }
    add_resource_box_section(doc, data, style)
    assert len(doc.tables) == 1
    table = doc.tables[0]
    # Header + 2 data rows = 3
    assert len(table.rows) == 3
    assert "Doc A" in table.cell(1, 0).text
