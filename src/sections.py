"""Section builder functions – one per layout type for Word documents.

Each builder receives a ``Document``, section data, and a :class:`Style`.

Mapping from presentation slide types to document sections:
- title       → Cover page (title + subtitle)
- content     → Heading + bullet list
- section-header → Section break with heading
- two-column  → Two-column table
- resource-box → Labelled resource table
"""

from __future__ import annotations

import os
from typing import TYPE_CHECKING

from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

if TYPE_CHECKING:
    from docx import Document
    from .style import Style


# ---------------------------------------------------------------------------
# Shared helpers
# ---------------------------------------------------------------------------


def _hex_to_rgb(hex_str: str) -> RGBColor:
    """Convert '#RRGGBB' to an RGBColor."""
    h = hex_str.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _add_image(doc: Document, img: dict) -> None:
    """Add an image to the document from a parsed ``**Image**`` dict."""
    path = img["path"]
    if not os.path.isfile(path):
        print(f"Warning: image not found: {path}, skipping.")
        return
    width = Inches(img.get("width", 4.0))
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(path, width=width)


def _set_cell_shading(cell, color_hex: str) -> None:
    """Set background shading on a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): color_hex.lstrip("#"),
    })
    shading.append(shd)


def _add_notes_block(doc: Document, notes: str, style: Style) -> None:
    """Optionally add notes as an indented italic block."""
    if not notes:
        return
    # Strip supplemental enrichment markers for clean output
    clean = notes.split("--- Supplemental")[0].strip()
    if not clean:
        return
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    run = p.add_run(clean)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)


# ---------------------------------------------------------------------------
# Section builders
# ---------------------------------------------------------------------------


def add_title_section(
    doc: Document,
    section_data: dict,
    style: Style,
) -> None:
    """Cover page: large centred title + subtitle."""
    # Add some spacing before title
    for _ in range(4):
        doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(section_data["title"])
    run.bold = True
    run.font.size = style.title_font
    run.font.color.rgb = _hex_to_rgb(style.title_color)

    subtitle = section_data.get("subtitle", "")
    if subtitle:
        # Handle <br> line breaks
        parts = [p.strip() for p in subtitle.split("<br>")]
        sub_p = doc.add_paragraph()
        sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for i, part in enumerate(parts):
            if i > 0:
                sub_p.add_run("\n")
            run = sub_p.add_run(part)
            run.font.size = style.subtitle_font
            run.font.color.rgb = _hex_to_rgb(style.subtitle_color)

    if section_data.get("image"):
        _add_image(doc, section_data["image"])

    _add_notes_block(doc, section_data.get("notes", ""), style)

    # Page break after title page
    doc.add_page_break()


def add_content_section(
    doc: Document,
    section_data: dict,
    style: Style,
) -> None:
    """Heading + bullet list."""
    heading = doc.add_heading(section_data["title"], level=2)
    for run in heading.runs:
        run.font.color.rgb = _hex_to_rgb(style.heading_color)

    for bullet in section_data.get("bullets", []):
        p = doc.add_paragraph(bullet, style="List Bullet")
        for run in p.runs:
            run.font.size = style.body_font

    if section_data.get("image"):
        _add_image(doc, section_data["image"])

    _add_notes_block(doc, section_data.get("notes", ""), style)


def add_section_header_section(
    doc: Document,
    section_data: dict,
    style: Style,
) -> None:
    """Section break with a prominent heading."""
    doc.add_page_break()

    # Add spacing
    for _ in range(3):
        doc.add_paragraph()

    heading = doc.add_heading(section_data["title"], level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.font.color.rgb = _hex_to_rgb(style.heading_color)
        run.font.size = style.title_font

    subtitle = section_data.get("subtitle", "")
    if subtitle:
        sub_p = doc.add_paragraph()
        sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sub_p.add_run(subtitle)
        run.font.size = style.subtitle_font
        run.font.color.rgb = _hex_to_rgb(style.subtitle_color)

    if section_data.get("image"):
        _add_image(doc, section_data["image"])

    _add_notes_block(doc, section_data.get("notes", ""), style)

    doc.add_page_break()


def add_two_column_section(
    doc: Document,
    section_data: dict,
    style: Style,
) -> None:
    """Two-column layout rendered as a two-column table."""
    heading = doc.add_heading(section_data["title"], level=2)
    for run in heading.runs:
        run.font.color.rgb = _hex_to_rgb(style.heading_color)

    left_bullets = section_data.get("left_bullets", [])
    right_bullets = section_data.get("right_bullets", [])
    max_rows = max(len(left_bullets), len(right_bullets), 1)

    table = doc.add_table(rows=max_rows, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Set column widths
    for row in table.rows:
        row.cells[0].width = Inches(3.25)
        row.cells[1].width = Inches(3.25)

    for i in range(max_rows):
        left_cell = table.cell(i, 0)
        right_cell = table.cell(i, 1)

        if i < len(left_bullets):
            p = left_cell.paragraphs[0]
            p.text = f"• {left_bullets[i]}"
            for run in p.runs:
                run.font.size = style.col_body_font

        if i < len(right_bullets):
            p = right_cell.paragraphs[0]
            p.text = f"• {right_bullets[i]}"
            for run in p.runs:
                run.font.size = style.col_body_font

    if section_data.get("image"):
        _add_image(doc, section_data["image"])

    _add_notes_block(doc, section_data.get("notes", ""), style)


def add_resource_box_section(
    doc: Document,
    section_data: dict,
    style: Style,
) -> None:
    """Resource boxes rendered as labelled tables with name/URL rows."""
    heading = doc.add_heading(section_data["title"], level=2)
    for run in heading.runs:
        run.font.color.rgb = _hex_to_rgb(style.heading_color)

    subtitle = section_data.get("subtitle", "")
    if subtitle:
        sub_p = doc.add_paragraph()
        run = sub_p.add_run(subtitle)
        run.font.size = style.subtitle_font
        run.font.color.rgb = _hex_to_rgb(style.accent_color)

    for box_data in section_data.get("boxes", []):
        label = box_data["label"]
        rows = box_data["rows"]

        # Label heading
        label_p = doc.add_paragraph()
        label_p.space_before = Pt(12)
        run = label_p.add_run(f"  {label}  ")
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = _hex_to_rgb(style.resource_label_color)

        if rows:
            table = doc.add_table(rows=len(rows) + 1, cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = True

            # Header row
            hdr = table.rows[0]
            hdr.cells[0].text = "Name"
            hdr.cells[1].text = "URL"
            for cell in hdr.cells:
                _set_cell_shading(cell, style.table_header_bg)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True
                        run.font.size = Pt(10)
                        run.font.color.rgb = _hex_to_rgb(style.table_header_color)

            # Data rows
            for i, row in enumerate(rows):
                data_row = table.rows[i + 1]
                data_row.cells[0].text = row["name"]
                data_row.cells[1].text = row.get("url", "")
                for p in data_row.cells[1].paragraphs:
                    for run in p.runs:
                        run.font.color.rgb = _hex_to_rgb(style.accent_color)
                        run.font.size = Pt(10)

        doc.add_paragraph()  # spacing

    if section_data.get("image"):
        _add_image(doc, section_data["image"])

    _add_notes_block(doc, section_data.get("notes", ""), style)


# ---------------------------------------------------------------------------
# Builder registry
# ---------------------------------------------------------------------------

SECTION_BUILDERS = {
    "title": add_title_section,
    "content": add_content_section,
    "section-header": add_section_header_section,
    "two-column": add_two_column_section,
    "resource-box": add_resource_box_section,
}
